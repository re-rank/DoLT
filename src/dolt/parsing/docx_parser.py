"""python-docx 기반 DOCX 파서."""

from __future__ import annotations

from dolt.errors import CorruptedFileError
from dolt.models.document import StructuredDocument
from dolt.models.section import Page, Section, Table
from dolt.parsing.base import BaseParser
from dolt.utils.logging import get_logger

logger = get_logger("parsing.docx")

# python-docx 헤딩 스타일 → 레벨 매핑
_HEADING_STYLES = {
    "Heading 1": 1, "Heading 2": 2, "Heading 3": 3,
    "Heading 4": 4, "Heading 5": 5, "Heading 6": 6,
    "heading 1": 1, "heading 2": 2, "heading 3": 3,
}


class DOCXParser(BaseParser):
    def supported_extensions(self) -> list[str]:
        return [".docx"]

    def parse(self, file_path: str, doc_id: str) -> StructuredDocument:
        from docx import Document as DocxDocument

        try:
            docx = DocxDocument(file_path)
        except Exception as e:
            raise CorruptedFileError(file_path, str(e)) from e

        sections: list[Section] = []
        tables: list[Table] = []
        text_parts: list[str] = []
        offset = 0
        section_counter = 0
        current_section_parts: list[str] = []

        for para in docx.paragraphs:
            style_name = para.style.name if para.style else ""
            level = _HEADING_STYLES.get(style_name)

            if level is not None:
                # 이전 섹션 본문 마무리
                if sections and current_section_parts:
                    sections[-1].content = "\n".join(current_section_parts).strip()
                current_section_parts = []

                section_counter += 1
                title = para.text.strip()
                sec_start = offset
                sections.append(
                    Section(
                        section_id=f"sec-{section_counter:03d}",
                        title=title,
                        level=level,
                        content="",
                        start_offset=sec_start,
                    )
                )
            else:
                if para.text.strip():
                    current_section_parts.append(para.text)

            text_parts.append(para.text)
            offset += len(para.text) + 1  # +1 for newline

        # 마지막 섹션 본문 마무리
        if sections and current_section_parts:
            sections[-1].content = "\n".join(current_section_parts).strip()

        # 섹션 end_offset 설정
        for i, sec in enumerate(sections):
            if i + 1 < len(sections):
                sec.end_offset = sections[i + 1].start_offset
            else:
                sec.end_offset = offset

        # 표 추출
        for t_idx, table in enumerate(docx.tables):
            rows_data: list[list[str]] = []
            for row in table.rows:
                rows_data.append([cell.text.strip() for cell in row.cells])

            if not rows_data:
                continue

            headers = rows_data[0]
            data_rows = rows_data[1:]
            md = _to_markdown_table(headers, data_rows)

            tables.append(
                Table(
                    table_id=f"tbl-{t_idx + 1}",
                    headers=headers,
                    rows=data_rows,
                    markdown=md,
                )
            )

        raw_text = "\n".join(text_parts)

        # 문서 메타데이터
        props = docx.core_properties
        doc_metadata = {
            k: v for k, v in {
                "title": props.title,
                "author": props.author,
                "subject": props.subject,
            }.items() if v
        }

        return StructuredDocument(
            doc_id=doc_id,
            source=file_path,
            raw_text=raw_text,
            pages=[
                Page(
                    page_number=1,
                    text=raw_text,
                    start_offset=0,
                    end_offset=len(raw_text),
                )
            ],
            sections=sections,
            tables=tables,
            metadata=doc_metadata,
            total_chars=len(raw_text),
            total_pages=1,
        )


def _to_markdown_table(headers: list[str], rows: list[list[str]]) -> str:
    if not headers:
        return ""
    lines = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join("---" for _ in headers) + " |",
    ]
    for row in rows:
        padded = row + [""] * (len(headers) - len(row))
        lines.append("| " + " | ".join(padded[: len(headers)]) + " |")
    return "\n".join(lines)
